"""
Multimodal Parser for AsterScope.

Capable of extracting text from raw strings, PDFs (with table detection),
DOCX (with table extraction), and Images (OCR).
"""
import io
import logging
from typing import Optional, List

try:
    from PIL import Image
except ImportError:
    pass

try:
    import pytesseract
except ImportError:
    pass

try:
    import pypdf
except ImportError:
    pass

try:
    import pdfplumber
except ImportError:
    pass

try:
    import docx
except ImportError:
    pass

logger = logging.getLogger(__name__)

class MultimodalParser:
    """
    Unified extraction class to parse text from various file formats,
    including structured table extraction from PDFs and DOCX files.
    """
    
    def __init__(self):
        # Check if tesseract is installed on the system
        self.tesseract_available = False
        try:
            if 'pytesseract' in globals():
                pytesseract.get_tesseract_version()
                self.tesseract_available = True
                logger.info("Tesseract OS binary verified. Image OCR enabled.")
        except Exception as e:
            logger.warning("Tesseract OCR not found on host OS. Image parsing will fallback to empty strings. (%s)", str(e))

        # Check for pdfplumber (table-aware PDF parsing)
        self.pdfplumber_available = 'pdfplumber' in globals()
        if self.pdfplumber_available:
            logger.info("pdfplumber available. Table-aware PDF extraction enabled.")

    def parse(self, file_bytes: bytes, mime_type: str) -> str:
        """
        Routes the file bytes to the appropriate parsing engine based on mime type.
        """
        if not file_bytes:
            return ""
            
        mime_type = mime_type.lower()
        
        # 1. Image OCR (PNG, JPG, TIFF)
        if mime_type.startswith("image/"):
            return self._parse_image(file_bytes)
            
        # 2. PDF Parsing (with table awareness)
        if mime_type == "application/pdf":
            return self._parse_pdf(file_bytes)
            
        # 3. DOCX Parsing (with table extraction)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._parse_docx(file_bytes)
            
        # 4. Raw Text/Markdown
        if mime_type in ["text/plain", "text/markdown", "text/csv"]:
            return file_bytes.decode("utf-8", errors="ignore")
            
        # Fallback
        logger.warning(f"Unsupported MIME type: {mime_type}. Attempting default raw string decode.")
        return file_bytes.decode("utf-8", errors="ignore")

    def _parse_image(self, file_bytes: bytes) -> str:
        """Extracts text from an image using Tesseract OCR."""
        if not self.tesseract_available or 'Image' not in globals() or 'pytesseract' not in globals():
            logger.warning("Image parsing requested but OCR dependencies are missing.")
            return ""
            
        try:
            image = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"OCR Parsing failed: {e}")
            return ""

    def _table_to_markdown(self, table: List[List[Optional[str]]]) -> str:
        """
        Converts a 2D table (list of rows) into a Markdown grid.
        Preserves structural semantics for downstream embedding.
        """
        if not table or len(table) < 1:
            return ""
        
        # Clean cells
        cleaned = []
        for row in table:
            cleaned.append([str(cell).strip() if cell else "" for cell in row])
        
        if not cleaned:
            return ""
        
        # Build markdown table
        header = cleaned[0]
        md_lines = ["| " + " | ".join(header) + " |"]
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        for row in cleaned[1:]:
            # Pad row to match header length
            padded = row + [""] * (len(header) - len(row))
            md_lines.append("| " + " | ".join(padded[:len(header)]) + " |")
        
        return "\n".join(md_lines)
            
    def _parse_pdf(self, file_bytes: bytes) -> str:
        """
        Extracts text AND tables from PDF layers.
        Uses pdfplumber for table-aware extraction when available,
        falls back to pypdf for text-only extraction.
        """
        sections = []
        
        # Try pdfplumber first for table-aware extraction
        if self.pdfplumber_available:
            try:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Extract tables from this page
                        tables = page.extract_tables()
                        table_regions = set()
                        
                        if tables:
                            for table in tables:
                                md_table = self._table_to_markdown(table)
                                if md_table:
                                    sections.append(f"[Table from page {page_num + 1}]\n{md_table}")
                        
                        # Extract remaining text (non-table content)
                        page_text = page.extract_text()
                        if page_text:
                            sections.append(page_text.strip())
                            
                return "\n\n".join(sections).strip()
            except Exception as e:
                logger.warning(f"pdfplumber table extraction failed, falling back to pypdf: {e}")
        
        # Fallback to pypdf (text-only)
        if 'pypdf' not in globals():
            logger.warning("PDF parsing requested but neither pdfplumber nor pypdf is installed.")
            return ""
            
        try:
            pdf = pypdf.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"PDF Parsing failed: {e}")
            return ""

    def _parse_docx(self, file_bytes: bytes) -> str:
        """
        Extracts text AND tables from DOCX files.
        Tables are converted to Markdown grids to preserve structure.
        """
        if 'docx' not in globals():
            logger.warning("DOCX parsing requested but python-docx is not installed.")
            return ""
            
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            sections = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    sections.append(para.text.strip())
            
            # Extract tables as Markdown grids
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                
                md_table = self._table_to_markdown(rows)
                if md_table:
                    sections.append(f"[Table {i + 1}]\n{md_table}")
            
            return "\n\n".join(sections).strip()
        except Exception as e:
            logger.error(f"DOCX Parsing failed: {e}")
            return ""
